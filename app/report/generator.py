import logging
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from jinja2 import Template
from weasyprint import HTML as WeasyPrintHTML  # noqa: N811

from app.models.bi import BIReport, BusinessProfile, MarketingAnalysis

logger = logging.getLogger(__name__)

HTML_TEMPLATE = Template("""
<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <style>
    @page { margin: 2cm; }
    body { font-family: 'Helvetica Neue', Arial, sans-serif; color: #333; line-height: 1.6; font-size: 11pt; }
    h1 { color: #1a237e; font-size: 22pt; border-bottom: 3px solid #1a237e; padding-bottom: 8px; }
    h2 { color: #283593; font-size: 16pt; margin-top: 24px; border-bottom: 1px solid #ddd; padding-bottom: 4px; }
    h3 { color: #3949ab; font-size: 13pt; margin-top: 16px; }
    .exec-summary { background: #e8eaf6; padding: 16px; border-radius: 6px; margin: 16px 0; }
    .exec-summary p { margin: 4px 0; }
    ul { padding-left: 20px; }
    li { margin: 4px 0; }
    .section { margin: 20px 0; }
    .label { font-weight: 600; color: #1a237e; }
    .footer {
      margin-top: 40px; padding-top: 16px; border-top: 1px solid #ddd;
      font-size: 9pt; color: #888; text-align: center;
    }
    table { width: 100%; border-collapse: collapse; margin: 12px 0; }
    th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
    th { background: #e8eaf6; font-weight: 600; }
  </style>
</head>
<body>
  <h1>Business Intelligence Report</h1>
  <p style="color: #666; font-size: 10pt;">{{ company }} &mdash; Generated {{ date }}</p>

  <div class="exec-summary">
    <h2>Executive Summary</h2>
    <p>{{ executive_summary }}</p>
  </div>

  <div class="section">
    <h2>1. Business Profile</h2>
    <p><span class="label">Company:</span> {{ profile.company_name }}</p>
    <p><span class="label">Tagline:</span> {{ profile.tagline }}</p>
    <p><span class="label">Industry:</span> {{ profile.industry }}</p>
    <p><span class="label">Value Proposition:</span> {{ profile.value_proposition }}</p>
    <h3>Offerings</h3>
    <ul>
    {% for o in profile.offerings %}
      <li>{{ o }}</li>
    {% endfor %}
    </ul>
    <h3>Target Audience</h3>
    <ul>
    {% for a in profile.target_audience %}
      <li>{{ a }}</li>
    {% endfor %}
    </ul>
  </div>

  <div class="section">
    <h2>2. Market Analysis</h2>
    <p><span class="label">Target Market:</span> {{ market.target_market }}</p>
    <p><span class="label">Market Size:</span> {{ market.market_size }}</p>
    <p><span class="label">Pricing Strategy:</span> {{ market.pricing_strategy }}</p>
    <p><span class="label">Market Positioning:</span> {{ market.market_positioning }}</p>
    <h3>Market Trends</h3>
    <ul>
    {% for t in market.market_trends %}
      <li>{{ t }}</li>
    {% endfor %}
    </ul>
  </div>

  <div class="section">
    <h2>3. Competitive Analysis</h2>
    <h3>Direct Competitors</h3>
    <ul>
    {% for c in competitive.direct_competitors %}
      <li><strong>{{ c.name if c.name else c }}</strong>: {{ c.description if c.description else '' }}</li>
    {% endfor %}
    </ul>
    <h3>Competitive Advantages</h3>
    <ul>
    {% for a in competitive.competitive_advantages %}
      <li>{{ a }}</li>
    {% endfor %}
    </ul>
    <p><span class="label">Threat Assessment:</span> {{ competitive.threat_assessment }}</p>
  </div>

  <div class="section">
    <h2>4. SWOT Analysis</h2>
    <table>
      <tr><th style="width:50%">Strengths</th><th style="width:50%">Weaknesses</th></tr>
      <tr>
        <td><ul>{% for s in swot.strengths %}<li>{{ s }}</li>{% endfor %}</ul></td>
        <td><ul>{% for w in swot.weaknesses %}<li>{{ w }}</li>{% endfor %}</ul></td>
      </tr>
      <tr><th>Opportunities</th><th>Threats</th></tr>
      <tr>
        <td><ul>{% for o in swot.opportunities %}<li>{{ o }}</li>{% endfor %}</ul></td>
        <td><ul>{% for t in swot.threats %}<li>{{ t }}</li>{% endfor %}</ul></td>
      </tr>
    </table>
  </div>

  <div class="section">
    <h2>5. Recommendations</h2>
    <ul>
    {% for r in recommendations %}
      <li>{{ r }}</li>
    {% endfor %}
    </ul>
  </div>

  <div class="section">
    <h2>6. Risk Factors</h2>
    <ul>
    {% for r in risk_factors %}
      <li>{{ r }}</li>
    {% endfor %}
    </ul>
  </div>

  {% if marketing %}
  <div class="section">
    <h2>7. Target Audience & Personas</h2>
    {% for p in marketing.personas %}
    <h3>{{ p.name or 'Persona' }}</h3>
    <p>{{ p.description }}</p>
    {% if p.pain_points %}
    <p><span class="label">Pain Points:</span></p>
    <ul>{% for pp in p.pain_points %}<li>{{ pp }}</li>{% endfor %}</ul>
    {% endif %}
    {% if p.goals %}
    <p><span class="label">Goals:</span></p>
    <ul>{% for g in p.goals %}<li>{{ g }}</li>{% endfor %}</ul>
    {% endif %}
    {% if p.channels %}
    <p><span class="label">Channels:</span></p>
    <ul>{% for c in p.channels %}<li>{{ c }}</li>{% endfor %}</ul>
    {% endif %}
    {% endfor %}
  </div>

  <div class="section">
    <h2>8. Brand Personality</h2>
    {% if marketing.brand_personality %}
    <p><span class="label">Archetype:</span> {{ marketing.brand_personality.archetype }}</p>
    <p><span class="label">Voice:</span> {{ marketing.brand_personality.voice_characteristics }}</p>
    <h3>Traits</h3>
    <ul>{% for t in marketing.brand_personality.traits %}<li>{{ t }}</li>{% endfor %}</ul>
    {% endif %}
  </div>

  <div class="section">
    <h2>9. People Also Ask</h2>
    {% for item in marketing.faq %}
    <p><strong>{{ item.question }}</strong></p>
    <p>{{ item.answer }}</p>
    {% endfor %}
  </div>

  <div class="section">
    <h2>10. Customer Journey</h2>
    {% for stage in marketing.customer_journey %}
    <h3>{{ stage.stage }}</h3>
    <p>{{ stage.description }}</p>
    {% if stage.touchpoints %}
    <p><span class="label">Touchpoints:</span></p>
    <ul>{% for t in stage.touchpoints %}<li>{{ t }}</li>{% endfor %}</ul>
    {% endif %}
    {% endfor %}
  </div>

  <div class="section">
    <h2>11. E-E-A-T &amp; GEO Optimization</h2>
    {% if marketing.eeat %}
    <h3>E-E-A-T Signals</h3>
    {% if marketing.eeat.experience %}
    <p><span class="label">Experience:</span></p>
    <ul>{% for s in marketing.eeat.experience %}<li>{{ s }}</li>{% endfor %}</ul>
    {% endif %}
    {% if marketing.eeat.expertise %}
    <p><span class="label">Expertise:</span></p>
    <ul>{% for s in marketing.eeat.expertise %}<li>{{ s }}</li>{% endfor %}</ul>
    {% endif %}
    {% if marketing.eeat.authoritativeness %}
    <p><span class="label">Authoritativeness:</span></p>
    <ul>{% for s in marketing.eeat.authoritativeness %}<li>{{ s }}</li>{% endfor %}</ul>
    {% endif %}
    {% if marketing.eeat.trustworthiness %}
    <p><span class="label">Trustworthiness:</span></p>
    <ul>{% for s in marketing.eeat.trustworthiness %}<li>{{ s }}</li>{% endfor %}</ul>
    {% endif %}
    {% endif %}
    {% if marketing.geo_tactics %}
    <h3>GEO Tactics</h3>
    <ul>
    {% for g in marketing.geo_tactics %}
    <li><strong>{{ g.tactic }}:</strong> {{ g.description }}</li>
    {% endfor %}
    </ul>
    {% endif %}
  </div>

  <div class="section">
    <h2>12. Call to Action</h2>
    <p>{{ marketing.call_to_action }}</p>
  </div>
  {% endif %}

  <div class="footer">
    <p>Business Intelligence System — Automated Analysis Report</p>
    <p>Source: {{ profile.website or 'N/A' }}</p>
  </div>
</body>
</html>
""")


def _sanitize_filename(name: str) -> str:
    result = re.sub(r"[^\w\s-]", "", name).strip()
    result = re.sub(r"[-\s]+", "_", result)
    return result or "report"


class ReportGenerator:
    def __init__(self, output_dir: str = "reports") -> None:
        self._output_dir = Path(output_dir)
        self._output_dir.mkdir(parents=True, exist_ok=True)

    def generate_pdf(
        self,
        report: BIReport,
        profile: BusinessProfile,
        company: str = "",
        marketing: MarketingAnalysis | None = None,
    ) -> Path:
        date_str = datetime.now().strftime("%B %d, %Y")
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")

        html = HTML_TEMPLATE.render(
            company=company or profile.company_name,
            date=date_str,
            executive_summary=report.executive_summary,
            profile=profile,
            market=report.market_analysis,
            competitive=report.competitive_analysis,
            swot=report.swot_analysis,
            recommendations=report.recommendations,
            risk_factors=report.risk_factors,
            marketing=marketing,
        )

        safe_name = _sanitize_filename(company or profile.company_name)
        filename = f"{safe_name}_{ts}.pdf"
        filepath = self._output_dir / filename
        WeasyPrintHTML(string=html).write_pdf(str(filepath))
        logger.info("PDF generated: %s", filepath)
        return filepath

    def generate_docx(
        self,
        report: BIReport,
        profile: BusinessProfile,
        company: str = "",
        marketing: MarketingAnalysis | None = None,
    ) -> Path:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Helvetica"

        doc.add_heading("Business Intelligence Report", level=1)
        p = doc.add_paragraph(
            f"{company or profile.company_name} — Generated "
            f"{datetime.now().strftime('%B %d, %Y')}"
        )
        p.italic = True

        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(report.executive_summary)

        doc.add_heading("1. Business Profile", level=2)
        for label, value in [
            ("Company", profile.company_name),
            ("Tagline", profile.tagline),
            ("Industry", profile.industry),
            ("Value Proposition", profile.value_proposition),
        ]:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(value)

        doc.add_heading("Offerings", level=3)
        for o in profile.offerings:
            doc.add_paragraph(o, style="List Bullet")

        doc.add_heading("Target Audience", level=3)
        for a in profile.target_audience:
            doc.add_paragraph(a, style="List Bullet")

        doc.add_heading("2. Market Analysis", level=2)
        for label, value in [
            ("Target Market", report.market_analysis.target_market),
            ("Market Size", report.market_analysis.market_size),
            ("Pricing Strategy", report.market_analysis.pricing_strategy),
            ("Market Positioning", report.market_analysis.market_positioning),
        ]:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(value)

        doc.add_heading("Market Trends", level=3)
        for t in report.market_analysis.market_trends:
            doc.add_paragraph(t, style="List Bullet")

        doc.add_heading("3. Competitive Analysis", level=2)
        doc.add_heading("Direct Competitors", level=3)
        for c in report.competitive_analysis.direct_competitors:
            name = c.get("name", str(c)) if isinstance(c, dict) else str(c)
            desc = c.get("description", "") if isinstance(c, dict) else ""
            doc.add_paragraph(f"{name}: {desc}", style="List Bullet")

        doc.add_heading("Competitive Advantages", level=3)
        for a in report.competitive_analysis.competitive_advantages:
            doc.add_paragraph(a, style="List Bullet")

        doc.add_heading("4. SWOT Analysis", level=2)
        for category, items in [
            ("Strengths", report.swot_analysis.strengths),
            ("Weaknesses", report.swot_analysis.weaknesses),
            ("Opportunities", report.swot_analysis.opportunities),
            ("Threats", report.swot_analysis.threats),
        ]:
            doc.add_heading(category, level=3)
            for item in items:
                doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("5. Recommendations", level=2)
        for r in report.recommendations:
            doc.add_paragraph(r, style="List Bullet")

        doc.add_heading("6. Risk Factors", level=2)
        for r in report.risk_factors:
            doc.add_paragraph(r, style="List Bullet")

        if marketing:
            doc.add_heading("7. Target Audience & Personas", level=2)
            for p in marketing.personas:
                doc.add_heading(p.name or "Persona", level=3)
                doc.add_paragraph(p.description)
                if p.pain_points:
                    doc.add_heading("Pain Points", level=4)
                    for pp in p.pain_points:
                        doc.add_paragraph(pp, style="List Bullet")
                if p.goals:
                    doc.add_heading("Goals", level=4)
                    for g in p.goals:
                        doc.add_paragraph(g, style="List Bullet")
                if p.channels:
                    doc.add_heading("Channels", level=4)
                    for c in p.channels:
                        doc.add_paragraph(c, style="List Bullet")

            doc.add_heading("8. Brand Personality", level=2)
            if marketing.brand_personality:
                for label, value in [
                    ("Archetype", marketing.brand_personality.archetype),
                    ("Voice", marketing.brand_personality.voice_characteristics),
                ]:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{label}: ")
                    run.bold = True
                    p.add_run(value)
                doc.add_heading("Traits", level=3)
                for t in marketing.brand_personality.traits:
                    doc.add_paragraph(t, style="List Bullet")

            doc.add_heading("9. People Also Ask", level=2)
            for item in marketing.faq:
                p = doc.add_paragraph()
                run = p.add_run(item.question)
                run.bold = True
                doc.add_paragraph(item.answer)

            doc.add_heading("10. Customer Journey", level=2)
            for stage in marketing.customer_journey:
                doc.add_heading(stage.stage, level=3)
                doc.add_paragraph(stage.description)
                if stage.touchpoints:
                    doc.add_heading("Touchpoints", level=4)
                    for t in stage.touchpoints:
                        doc.add_paragraph(t, style="List Bullet")

            doc.add_heading("11. E-E-A-T & GEO Optimization", level=2)
            if marketing.eeat:
                doc.add_heading("E-E-A-T Signals", level=3)
                for label, items in [
                    ("Experience", marketing.eeat.experience),
                    ("Expertise", marketing.eeat.expertise),
                    ("Authoritativeness", marketing.eeat.authoritativeness),
                    ("Trustworthiness", marketing.eeat.trustworthiness),
                ]:
                    if items:
                        doc.add_heading(label, level=4)
                        for s in items:
                            doc.add_paragraph(s, style="List Bullet")
            if marketing.geo_tactics:
                doc.add_heading("GEO Tactics", level=3)
                for g in marketing.geo_tactics:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{g.tactic}: ")
                    run.bold = True
                    p.add_run(g.description)

            doc.add_heading("12. Call to Action", level=2)
            doc.add_paragraph(marketing.call_to_action)

        safe_name = _sanitize_filename(company or profile.company_name)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe_name}_{ts}.docx"
        filepath = self._output_dir / filename
        doc.save(str(filepath))
        logger.info("DOCX generated: %s", filepath)
        return filepath
